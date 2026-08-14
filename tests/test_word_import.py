from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image

from examdesk.domain.enums import QuestionType, UsageScope
from examdesk.importers import parse_word_docx


def make_png(path: Path) -> None:
    image = Image.new("RGB", (80, 50), (20, 130, 90))
    image.save(path, format="PNG")


def test_word_import_reads_inline_metadata_and_question_image(tmp_path: Path) -> None:
    image_path = tmp_path / "题图.png"
    make_png(image_path)
    document = Document()
    document.add_paragraph("1. 临时用地复垦后应如何认定？")
    document.add_picture(str(image_path), width=Inches(1))
    document.add_paragraph("A. 直接变更")
    document.add_paragraph("B. 结合现状认定")
    document.add_paragraph("C. 保持原分类")
    document.add_paragraph("D. 删除记录")
    document.add_paragraph("题型：单选")
    document.add_paragraph("答案：B")
    document.add_paragraph("依据：培训手册第三条")
    document.add_paragraph("分值：2")
    document.add_paragraph("章节：安全规范")
    document.add_paragraph("适用年度：2026")
    document.add_paragraph("来源：2026年度培训手册")
    document.add_paragraph("使用范围：练习和考试")
    path = tmp_path / "inline.docx"
    document.save(path)

    preview = parse_word_docx(path)
    candidate = preview.candidates[0]

    assert preview.error_count == 0
    assert candidate.question.question_type is QuestionType.SINGLE
    assert candidate.question.correct_option_keys == {"B"}
    assert candidate.question.usage_scope is UsageScope.BOTH
    assert candidate.question.applicable_year == 2026
    assert len(candidate.images) == 1
    assert candidate.images[0].owner_key == "stem"


def test_word_import_maps_trailing_answer_list_and_infers_multiple_choice(tmp_path: Path) -> None:
    document = Document()
    for number, stem in ((1, "第一题"), (2, "第二题")):
        document.add_paragraph(f"{number}. {stem}")
        for key, text in zip("ABCD", ("甲", "乙", "丙", "丁"), strict=True):
            document.add_paragraph(f"{key}. {text}")
        document.add_paragraph("依据：培训手册")
    document.add_paragraph("答案")
    document.add_paragraph("1.A  2.BC")
    path = tmp_path / "trailing.docx"
    document.save(path)

    preview = parse_word_docx(path)

    assert len(preview.candidates) == 2
    assert preview.candidates[0].question.correct_option_keys == {"A"}
    assert preview.candidates[0].question.question_type is QuestionType.SINGLE
    assert preview.candidates[1].question.correct_option_keys == {"B", "C"}
    assert preview.candidates[1].question.question_type is QuestionType.MULTIPLE


def test_word_import_marks_inline_and_trailing_answer_conflict(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("1. 冲突题")
    document.add_paragraph("A. 甲")
    document.add_paragraph("B. 乙")
    document.add_paragraph("答案：A")
    document.add_paragraph("依据：依据")
    document.add_paragraph("答案")
    document.add_paragraph("1.B")
    path = tmp_path / "conflict.docx"
    document.save(path)

    preview = parse_word_docx(path)

    assert any(issue.code == "answer_conflict" for issue in preview.issues)


def test_word_import_reads_fill_question_with_repeated_score(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("1. 填写（1）（2）（3）")
    document.add_paragraph("题型：填空")
    document.add_paragraph("答案：@1-3|甲;乙;丙")
    document.add_paragraph("分值：0.5")
    document.add_paragraph("依据：依据")
    path = tmp_path / "fill.docx"
    document.save(path)

    preview = parse_word_docx(path)
    question = preview.candidates[0].question

    assert question.question_type is QuestionType.FILL
    assert str(question.score) == "1.5"
    assert question.unordered_groups[0].indexes == (1, 2, 3)

